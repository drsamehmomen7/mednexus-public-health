"""
Tests for document parsing (DOCX/TXT text extraction).
"""

from io import BytesIO

import pytest
from docx import Document

from app.services.document_parsing import (
    UnsupportedDocumentType,
    extract_text,
    extract_text_from_docx,
)


def _make_docx_bytes(paragraphs, table_rows=None):
    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table_rows:
        table = doc.add_table(rows=0, cols=len(table_rows[0]))
        for row_values in table_rows:
            row = table.add_row()
            for cell, value in zip(row.cells, row_values):
                cell.text = value
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_extracts_paragraph_text():
    file_bytes = _make_docx_bytes(["Line one.", "Line two."])
    text = extract_text_from_docx(file_bytes)
    assert "Line one." in text
    assert "Line two." in text


def test_extracts_table_text():
    file_bytes = _make_docx_bytes(
        ["Header"],
        table_rows=[["Facility", "Ardiya Clinic"], ["Region", "Farwaniya"]],
    )
    text = extract_text_from_docx(file_bytes)
    assert "Facility" in text
    assert "Ardiya Clinic" in text
    assert "Farwaniya" in text


def test_empty_paragraphs_are_skipped():
    file_bytes = _make_docx_bytes(["Real line.", "", "   ", "Another real line."])
    text = extract_text_from_docx(file_bytes)
    lines = [l for l in text.split("\n") if l]
    assert lines == ["Real line.", "Another real line."]


def test_extract_text_dispatches_docx_by_extension():
    file_bytes = _make_docx_bytes(["Some content."])
    assert extract_text("report.docx", file_bytes) == extract_text_from_docx(file_bytes)


def test_extract_text_dispatches_txt_by_extension():
    file_bytes = "Plain text content.".encode("utf-8")
    assert extract_text("report.txt", file_bytes) == "Plain text content."


def test_unsupported_extension_raises():
    with pytest.raises(UnsupportedDocumentType):
        extract_text("report.pdf", b"whatever")


def test_preserves_true_reading_order_across_paragraphs_and_tables():
    """
    Real bug found while testing generated case-report forms: grouping
    all paragraphs before all tables put a footer paragraph (written
    AFTER the table in the document) ahead of the table's field data in
    the extracted text — disorienting even though it happened not to
    break extraction in that case. A letterhead paragraph, then a
    table, then a footer paragraph must come out in that order.
    """
    doc = Document()
    doc.add_paragraph("LETTERHEAD")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    doc.add_paragraph("FOOTER")

    buf = BytesIO()
    doc.save(buf)
    text = extract_text_from_docx(buf.getvalue())

    lines = [l for l in text.split("\n") if l]
    assert lines == ["LETTERHEAD", "Field Value", "FOOTER"]
